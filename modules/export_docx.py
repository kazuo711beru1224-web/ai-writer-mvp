# modules/export_docx.py
from pathlib import Path
from datetime import datetime
from docx import Document

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)


def export_articles_to_docx(main_keyword: str, articles: list[str]) -> list[Path]:
    """
    複数の記事を Word（.docx）として outputs/ に保存する。
    返り値：保存した docx ファイルの Path のリスト
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_kw = main_keyword.replace(" ", "_").replace("　", "_")

    saved_paths = []

    for idx, article in enumerate(articles, start=1):
        filename = f"output_{timestamp}_{idx:02d}.docx"
        file_path = OUTPUT_DIR / filename

        doc = Document()
        doc.add_heading(f"{main_keyword}", level=1)
        doc.add_paragraph(article)

        doc.save(file_path)
        saved_paths.append(file_path)

    return saved_paths
