from docx import Document
from docx.shared import Pt
from datetime import datetime
import json

def build_docx(title, sections, out_path, styles):
    doc = Document()

    # タイトル
    style_title = styles.get("title", {})
    title_p = doc.add_paragraph(title)
    run = title_p.runs[0]
    run.bold = style_title.get("bold", True)
    run.font.size = Pt(style_title.get("font_size", 18))

    # セクション
    for s in sections:
        level = s.get("level", 2)
        heading_style = f"heading{level}" if f"heading{level}" in styles else "body"
        style = styles.get(heading_style, {})
        doc.add_paragraph(s["title"], style="Heading {}".format(level))
        p = doc.add_paragraph(s["body"])
        p.style.font.size = Pt(style.get("font_size", 11))

    doc.add_paragraph(f"\n生成日時: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.save(out_path)
    return out_path
